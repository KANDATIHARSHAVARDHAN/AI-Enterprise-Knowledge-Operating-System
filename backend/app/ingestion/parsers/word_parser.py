"""
EKOS Word Document Parser
Extracts text from .docx files using python-docx.
"""

from pathlib import Path
from docx import Document
from app.utils.logger import logger
from app.utils.exceptions import IngestionError


class WordParser:
    """Parse Word (.docx) files and extract text content."""

    SUPPORTED_EXTENSIONS = {".docx", ".doc"}

    def parse(self, file_path: str) -> list[dict]:
        """
        Parse a Word document and extract text.

        Returns:
            List with single dict containing full document text
        """
        path = Path(file_path)
        if not path.exists():
            raise IngestionError(f"File not found: {file_path}", filename=path.name)

        try:
            doc = Document(str(path))
            sections = []
            current_section = {"title": "Document Start", "content": []}

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Detect headings as section breaks
                if para.style and para.style.name.startswith("Heading"):
                    if current_section["content"]:
                        sections.append(current_section)
                    current_section = {"title": text, "content": []}
                else:
                    current_section["content"].append(text)

            # Add last section
            if current_section["content"]:
                sections.append(current_section)

            # Also extract tables
            table_texts = self._extract_tables(doc)

            # Build result pages (one per section)
            pages = []
            for i, section in enumerate(sections, 1):
                content = f"## {section['title']}\n\n" + "\n".join(section["content"])
                pages.append({
                    "content": content,
                    "metadata": {
                        "source": path.name,
                        "file_type": "docx",
                        "section_number": i,
                        "section_title": section["title"],
                        "file_path": str(path),
                    }
                })

            # Add tables as a separate section
            if table_texts:
                pages.append({
                    "content": "[Tables]\n" + table_texts,
                    "metadata": {
                        "source": path.name,
                        "file_type": "docx",
                        "section_number": len(sections) + 1,
                        "section_title": "Tables",
                        "file_path": str(path),
                    }
                })

            logger.info(f"Parsed Word doc: {path.name} ({len(pages)} sections)")
            return pages

        except Exception as e:
            raise IngestionError(f"Failed to parse Word doc: {e}", filename=path.name)

    def _extract_tables(self, doc: Document) -> str:
        """Extract tables from a Word document."""
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                table_texts.append("\n".join(rows))
        return "\n\n".join(table_texts)
